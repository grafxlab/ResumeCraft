from __future__ import annotations

import io
import json

from app.services import llm

PARSE_SYSTEM = (
    "You extract structured data from resumes. Return ONLY valid JSON — no "
    "prose, no markdown fences. Never invent information; use null or empty "
    "arrays when something is not present in the resume."
)
MAX_RESUME_TEXT_CHARS = 100_000

# The exact shape the frontend profile form expects.
_SCHEMA = """{
  "full_name": string,
  "email": string | null,
  "phone": string | null,
  "location": string | null,
  "summary": string | null,
  "skills": string[],
  "experience": [
    {"title": string, "company": string, "start": string | null,
     "end": string | null, "highlights": string[]}
  ],
  "education": [
        {"degree": string, "institution": string, "year": string | null,
         "location": string | null, "details": string | null}
  ],
  "links": { "label": "url" }
}"""


class ResumeParseError(RuntimeError):
    pass


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from a PDF, DOCX, or TXT resume upload."""
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if name.endswith(".docx"):
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(data))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        table_rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for table in doc.tables
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        ]
        return "\n".join([*paragraphs, *table_rows])

    if name.endswith((".txt", ".md", ".json")):
        return data.decode("utf-8", errors="ignore")

    raise ResumeParseError(
        "Unsupported file type. Please upload a PDF, DOCX, TXT, or JSON file."
    )


def _extract_json(text: str) -> dict:
    """Best-effort parse of a JSON object from an LLM response.

    Handles code fences and any leading/trailing prose by falling back to the
    substring between the first '{' and the last '}'.
    """
    cleaned = text.strip()

    # Strip a leading ```json / ``` fence and trailing ``` if present.
    if cleaned.startswith("```"):
        cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Fall back to the outermost object braces.
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start != -1 and end != -1 and end > start:
        return json.loads(cleaned[start : end + 1])

    raise json.JSONDecodeError("No JSON object found", cleaned, 0)


async def parse_resume(text: str, user_id: int | None = None) -> dict:
    """Use the LLM to convert raw resume text into structured profile data."""
    if not text.strip():
        raise ResumeParseError("No text could be extracted from the file.")

    prompt = (
        "Extract the resume below into JSON matching EXACTLY this schema. "
        "Read the complete document, including Education and other sections "
        "near the end. Preserve coursework and incomplete degree details. "
        "Respond with the JSON object only — no explanation, no code fences.\n"
        f"{_SCHEMA}\n\n"
        "RESUME TEXT:\n"
        f"{text[:MAX_RESUME_TEXT_CHARS]}"
    )
    raw = await llm.complete(
        PARSE_SYSTEM,
        prompt,
        max_tokens=4096,
        operation="resume_import",
        user_id=user_id,
    )

    try:
        return _extract_json(raw)
    except json.JSONDecodeError as exc:
        raise ResumeParseError(
            "The AI response could not be parsed as JSON. Try again."
        ) from exc
